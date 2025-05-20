import base64
import streamlit as st
import requests
from web3 import Web3
import time
from streamlit_pdf_viewer import pdf_viewer
import docx
import io
from docx import Document
from docx.shared import RGBColor

st.set_page_config(page_title="ScholarChain", layout="centered")
st.title("📚 ScholarChain – Academic Publishing Platform")

# ==== Blockchain Setup ====
PINATA_JWT = "eyJhbGciOiJIUzI1NiIsInR5cCI6IkpXVCJ9.eyJ1c2VySW5mb3JtYXRpb24iOnsiaWQiOiJkOTU2YmU5OS1kMjNjLTQ1NGMtOWY0Yy04NmM4ZGFiMjU1MDYiLCJlbWFpbCI6InNocmV5YS5iaGFub3QxM0BnbWFpbC5jb20iLCJlbWFpbF92ZXJpZmllZCI6dHJ1ZSwicGluX3BvbGljeSI6eyJyZWdpb25zIjpbeyJkZXNpcmVkUmVwbGljYXRpb25Db3VudCI6MSwiaWQiOiJGUkExIn0seyJkZXNpcmVkUmVwbGljYXRpb25Db3VudCI6MSwiaWQiOiJOWUMxIn1dLCJ2ZXJzaW9uIjoxfSwibWZhX2VuYWJsZWQiOmZhbHNlLCJzdGF0dXMiOiJBQ1RJVkUifSwiYXV0aGVudGljYXRpb25UeXBlIjoic2NvcGVkS2V5Iiwic2NvcGVkS2V5S2V5IjoiOGU3YWQ3NWI5NGQyZTZjOThmNjYiLCJzY29wZWRLZXlTZWNyZXQiOiJiZDQ4ZjI1ZTExYzc5OGZlOGUyOGRlM2VhYjkyNjcwMDRjMzhlMDRjZTBjZWEwM2E2ZmE4YjM5MzhkOGJhMWYxIiwiZXhwIjoxNzc4NTk5ODU4fQ.L-wSbPcRlKN_nzTIGMdZqv4FRzlbrnXJWuz7J1AZuc0"
CONTRACT_ADDRESS = "0xe7d16AE99BfD3C5d35c442b366A6B331AFFa8Bc9"
RPC_URL = "http://localhost:7545"
w3 = Web3(Web3.HTTPProvider(RPC_URL))

# ==== Load ABI (truncated for brevity) ====
CONTRACT_ABI = [
	{
		"inputs": [],
		"stateMutability": "nonpayable",
		"type": "constructor"
	},
	{
		"anonymous": False,
		"inputs": [
			{
				"indexed": True,
				"internalType": "uint256",
				"name": "paperId",
				"type": "uint256"
			},
			{
				"indexed": True,
				"internalType": "address",
				"name": "author",
				"type": "address"
			},
			{
				"indexed": False,
				"internalType": "string",
				"name": "title",
				"type": "string"
			},
			{
				"indexed": False,
				"internalType": "string",
				"name": "ipfsHash",
				"type": "string"
			}
		],
		"name": "PaperSubmitted",
		"type": "event"
	},
	{
		"anonymous": False,
		"inputs": [
			{
				"indexed": True,
				"internalType": "uint256",
				"name": "paperId",
				"type": "uint256"
			},
			{
				"indexed": True,
				"internalType": "address",
				"name": "reviewer",
				"type": "address"
			},
			{
				"indexed": False,
				"internalType": "string",
				"name": "comment",
				"type": "string"
			}
		],
		"name": "ReviewAdded",
		"type": "event"
	},
	{
		"anonymous": False,
		"inputs": [
			{
				"indexed": True,
				"internalType": "address",
				"name": "user",
				"type": "address"
			},
			{
				"indexed": False,
				"internalType": "string",
				"name": "name",
				"type": "string"
			},
			{
				"indexed": False,
				"internalType": "enum ResearchPortal.Role",
				"name": "role",
				"type": "uint8"
			}
		],
		"name": "UserRegistered",
		"type": "event"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "_paperId",
				"type": "uint256"
			},
			{
				"internalType": "string",
				"name": "comment",
				"type": "string"
			}
		],
		"name": "addReview",
		"outputs": [],
		"stateMutability": "nonpayable",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "_paperId",
				"type": "uint256"
			},
			{
				"internalType": "address",
				"name": "_user",
				"type": "address"
			}
		],
		"name": "getAccessLevel",
		"outputs": [
			{
				"internalType": "uint8",
				"name": "",
				"type": "uint8"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [],
		"name": "getAllPapers",
		"outputs": [
			{
				"components": [
					{
						"internalType": "uint256",
						"name": "paperId",
						"type": "uint256"
					},
					{
						"internalType": "address",
						"name": "author",
						"type": "address"
					},
					{
						"internalType": "string",
						"name": "title",
						"type": "string"
					},
					{
						"internalType": "string",
						"name": "ipfsHash",
						"type": "string"
					},
					{
						"internalType": "uint256",
						"name": "timestamp",
						"type": "uint256"
					},
					{
						"internalType": "string[]",
						"name": "reviews",
						"type": "string[]"
					}
				],
				"internalType": "struct ResearchPortal.Paper[]",
				"name": "",
				"type": "tuple[]"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "_paperId",
				"type": "uint256"
			}
		],
		"name": "getPaper",
		"outputs": [
			{
				"internalType": "uint256",
				"name": "",
				"type": "uint256"
			},
			{
				"internalType": "address",
				"name": "",
				"type": "address"
			},
			{
				"internalType": "string",
				"name": "",
				"type": "string"
			},
			{
				"internalType": "string",
				"name": "",
				"type": "string"
			},
			{
				"internalType": "uint256",
				"name": "",
				"type": "uint256"
			},
			{
				"internalType": "string[]",
				"name": "",
				"type": "string[]"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "_paperId",
				"type": "uint256"
			},
			{
				"internalType": "address",
				"name": "_user",
				"type": "address"
			},
			{
				"internalType": "uint8",
				"name": "_accessLevel",
				"type": "uint8"
			}
		],
		"name": "grantAccess",
		"outputs": [],
		"stateMutability": "nonpayable",
		"type": "function"
	},
	{
		"inputs": [],
		"name": "owner",
		"outputs": [
			{
				"internalType": "address",
				"name": "",
				"type": "address"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "",
				"type": "uint256"
			},
			{
				"internalType": "address",
				"name": "",
				"type": "address"
			}
		],
		"name": "paperAccess",
		"outputs": [
			{
				"internalType": "uint8",
				"name": "",
				"type": "uint8"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [],
		"name": "paperCount",
		"outputs": [
			{
				"internalType": "uint256",
				"name": "",
				"type": "uint256"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "uint256",
				"name": "",
				"type": "uint256"
			}
		],
		"name": "papers",
		"outputs": [
			{
				"internalType": "uint256",
				"name": "paperId",
				"type": "uint256"
			},
			{
				"internalType": "address",
				"name": "author",
				"type": "address"
			},
			{
				"internalType": "string",
				"name": "title",
				"type": "string"
			},
			{
				"internalType": "string",
				"name": "ipfsHash",
				"type": "string"
			},
			{
				"internalType": "uint256",
				"name": "timestamp",
				"type": "uint256"
			}
		],
		"stateMutability": "view",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "address",
				"name": "_user",
				"type": "address"
			},
			{
				"internalType": "string",
				"name": "_name",
				"type": "string"
			},
			{
				"internalType": "uint8",
				"name": "_role",
				"type": "uint8"
			}
		],
		"name": "register",
		"outputs": [],
		"stateMutability": "nonpayable",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "string",
				"name": "_title",
				"type": "string"
			},
			{
				"internalType": "string",
				"name": "_ipfsHash",
				"type": "string"
			}
		],
		"name": "submitPaper",
		"outputs": [],
		"stateMutability": "nonpayable",
		"type": "function"
	},
	{
		"inputs": [
			{
				"internalType": "address",
				"name": "",
				"type": "address"
			}
		],
		"name": "users",
		"outputs": [
			{
				"internalType": "enum ResearchPortal.Role",
				"name": "role",
				"type": "uint8"
			},
			{
				"internalType": "string",
				"name": "name",
				"type": "string"
			}
		],
		"stateMutability": "view",
		"type": "function"
	}
]

contract = w3.eth.contract(address=CONTRACT_ADDRESS, abi=CONTRACT_ABI)
def highlight_changes(original_text, edited_text):
    highlighted_doc = Document()
    original_lines = original_text.splitlines()
    edited_lines = edited_text.splitlines()
    for o, e in zip(original_lines, edited_lines):
        para = highlighted_doc.add_paragraph()
        if o != e:
            run = para.add_run(e)
            run.font.highlight_color = 3
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            para.add_run(e)
    buffer = io.BytesIO()
    highlighted_doc.save(buffer)
    buffer.seek(0)
    return buffer

def get_user_role(eth_address):
    try:
        role, name = contract.functions.users(Web3.to_checksum_address(eth_address)).call()
        return ("Author", name) if role == 1 else ("Faculty", name) if role == 2 else ("Reviewer", name) if role == 3 else ("Unregistered", name)
    except:
        return "Unregistered", ""

def get_events(event_name):
    event_filter = getattr(contract.events, event_name).create_filter(from_block=0, to_block='latest')
    return event_filter.get_all_entries()

def register_user(name, role, eth_address, private_key):
    try:
        role_map = {"Author": 1, "Faculty": 2, "Reviewer": 3}
        role_code = role_map.get(role, 0)
        tx = contract.functions.register(Web3.to_checksum_address(eth_address), name, role_code).build_transaction({
            'chainId': w3.eth.chain_id,
            'from': eth_address,
            'nonce': w3.eth.get_transaction_count(eth_address),
            'gas': 300000,
            'gasPrice': w3.to_wei('50', 'gwei')
        })
        signed_tx = w3.eth.account.sign_transaction(tx, private_key)
        tx_hash = w3.eth.send_raw_transaction(signed_tx.raw_transaction)
        receipt = w3.eth.wait_for_transaction_receipt(tx_hash, timeout=120)
        return receipt.status == 1
    except Exception as e:
        st.error(f"Registration failed: {str(e)}")
        return False

def upload_to_pinata(file_bytes, filename):
    url = "https://api.pinata.cloud/pinning/pinFileToIPFS"
    headers = {"Authorization": f"Bearer {PINATA_JWT}"}
    try:
        response = requests.post(url, files={"file": (filename, file_bytes)}, headers=headers, timeout=30)
        return response.json().get("IpfsHash") if response.status_code == 200 else None
    except requests.exceptions.RequestException as e:
        st.error(f"❌ Network error during upload: {str(e)}")
        return None

def fetch_ipfs_file(ipfs_hash):
    try:
        response = requests.get(f"https://gateway.pinata.cloud/ipfs/{ipfs_hash}")
        return response.content if response.status_code == 200 else None
    except:
        return None

def get_access_level(paper_id, user_address):
    try:
        return contract.functions.getAccessLevel(paper_id, Web3.to_checksum_address(user_address)).call()
    except:
        return 0  # No access
# ==== Sidebar: User Management ====
st.sidebar.header("👥 User Management")
if "users" not in st.session_state:
    st.session_state.users = []

with st.sidebar.expander("➕ Register New User"):
    reg_eth_address = st.text_input("Ethereum Address", key="reg_eth_address").strip()
    reg_private_key = st.text_input("Private Key", type="password", key="reg_private_key").strip()
    reg_name = st.text_input("Full Name", key="reg_name")
    reg_role = st.selectbox("Role", ["Author", "Faculty", "Reviewer"], key="reg_role")
    if st.button("Register", key="register_btn"):
        if not (reg_eth_address and reg_private_key and reg_name):
            st.warning("Please fill all fields.")
        elif register_user(reg_name, reg_role, reg_eth_address, reg_private_key):
            st.success(f"{reg_name} registered as {reg_role}.")
            st.session_state.users.append({
                "eth_address": reg_eth_address,
                "private_key": reg_private_key,
                "name": reg_name
            })

if st.session_state.users:
    user_options = [f"{u['name']} [{u['eth_address'][:6]}...]" for u in st.session_state.users]
    selected_idx = st.sidebar.selectbox(
        "Switch Active User",
        options=list(range(len(user_options))),
        format_func=lambda i: user_options[i],
        index=st.session_state.get("active_user_idx", 0),
        key="active_user_idx"
    )
    active_user = st.session_state.users[selected_idx]
    eth_address = active_user["eth_address"]
    private_key = active_user["private_key"]
    user_role, user_name = get_user_role(eth_address)
    st.sidebar.success(f"Active: {user_name} ({user_role})")
else:
    st.info("No users registered yet.")
    st.stop()
st.header("📚 Published Papers")
try:
    paper_count = contract.functions.paperCount().call()
    if paper_count == 0:
        st.info("No papers published yet.")
    else:
        for pid in range(1, paper_count + 1):
            paper = contract.functions.getPaper(pid).call()
            paper_id, author, title, ipfs_hash, timestamp, reviews = paper
            access_level = 3 if author.lower() == eth_address.lower() else get_access_level(paper_id, eth_address)

            if access_level > 0 or user_role in ["Faculty", "Reviewer"]:
                with st.expander(f"Paper #{paper_id}: {title}"):
                    st.markdown(f"""
**Author:** `{author}`  
**Title:** {title}  
**Submitted:** {time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(timestamp))}  
**Reviews/Comments:** {len(reviews)}  
**Access Level:** {['None', 'View', 'Comment', 'Edit'][access_level]}
                    """)
                    file_bytes = fetch_ipfs_file(ipfs_hash)
                    if file_bytes:
                        if ipfs_hash.endswith('.pdf') or file_bytes[:4] == b'%PDF':
                            st.markdown("**📄 PDF Preview:**")
                            pdf_base64 = base64.b64encode(file_bytes).decode('utf-8')
                            st.components.v1.html(f"""
                                <iframe src="data:application/pdf;base64,{pdf_base64}"
                                        width="100%" height="900px" style="border:none;"></iframe>
                            """, height=900)
                            pdf_viewer(file_bytes, width=1000, height=800, key=f"pdf_{paper_id}")
                            if access_level >= 1:
                                st.download_button("Download PDF", file_bytes, file_name=f"{title}.pdf", mime="application/pdf")

                        elif ipfs_hash.endswith('.docx') or file_bytes[:4] == b'PK\x03\x04':
                            st.markdown("**📄 DOCX Preview:**")
                            try:
                                doc = docx.Document(io.BytesIO(file_bytes))
                                original_text = "\n".join([para.text for para in doc.paragraphs])

                                if access_level == 3:
                                    edited_text = st.text_area("📝 Edit Document", original_text, height=600)
                                    if edited_text != original_text:
                                        buffer = highlight_changes(original_text, edited_text)
                                        st.download_button("💾 Download Edited with Highlights", buffer, file_name=f"{title}_edited.docx")
                                    else:
                                        st.info("✅ No changes made.")
                                elif access_level == 2:
                                    st.text_area("💬 Comment Mode", original_text, height=600, disabled=True)
                                else:
                                    st.text_area("👁️ View Only", original_text, height=600, disabled=True)

                            except Exception as e:
                                st.warning(f"Could not load DOCX: {str(e)}")

                    if reviews:
                        st.markdown("**📌 Reviews:**")
                        for idx, review in enumerate(reviews):
                            st.markdown(f"{idx + 1}. {review}")

                    if access_level >= 2:
                        review_input = st.text_area("🗣️ Add Review", key=f"review_{paper_id}")
                        if st.button(f"Submit Review #{paper_id}", key=f"btn_{paper_id}"):
                            if review_input.strip():
                                try:
                                    tx = contract.functions.addReview(int(paper_id), review_input).build_transaction({
                                        'from': eth_address,
                                        'nonce': w3.eth.get_transaction_count(eth_address),
                                        'gas': 300000,
                                        'gasPrice': w3.to_wei('50', 'gwei')
                                    })
                                    signed_tx = w3.eth.account.sign_transaction(tx, private_key)
                                    tx_hash = w3.eth.send_raw_transaction(signed_tx.raw_transaction)
                                    receipt = w3.eth.wait_for_transaction_receipt(tx_hash)
                                    st.success(f"✅ Review submitted! TX Hash: `{tx_hash.hex()}`")
                                except Exception as e:
                                    st.error(f"Review failed: {str(e)}")
                            else:
                                st.warning("Please write something first.")
except Exception as e:
    st.error(f"Error loading papers: {str(e)}")
# ==== Author Actions ====
if user_role == "Author":
    st.header("📝 Author Panel")
    st.subheader("📤 Submit a New Paper")

    title = st.text_input("Paper Title")
    file = st.file_uploader("Upload Paper (PDF/DOCX)", type=["pdf", "docx"])

    if title and file:
        ipfs_hash = upload_to_pinata(file.read(), file.name)
        if ipfs_hash and st.button("🚀 Submit to Blockchain"):
            try:
                tx = contract.functions.submitPaper(title, ipfs_hash).build_transaction({
                    'from': eth_address,
                    'nonce': w3.eth.get_transaction_count(eth_address),
                    'gas': 500000,
                    'gasPrice': w3.to_wei('50', 'gwei')
                })
                signed_tx = w3.eth.account.sign_transaction(tx, private_key)
                tx_hash = w3.eth.send_raw_transaction(signed_tx.raw_transaction)
                receipt = w3.eth.wait_for_transaction_receipt(tx_hash)
                if receipt.status == 1:
                    st.balloons()
                    st.success(f"✅ Paper Submitted! TX Hash: `{tx_hash.hex()}`")
                else:
                    st.error("❌ Transaction failed.")
            except Exception as e:
                st.error(f"Submission error: {str(e)}")

    # ==== Grant Access Section ====
    st.subheader("🔐 Grant Access to a Paper")
    try:
        total_papers = contract.functions.paperCount().call()
        my_paper_ids = [
            pid for pid in range(1, total_papers + 1)
            if contract.functions.getPaper(pid).call()[1].lower() == eth_address.lower()
        ]

        if my_paper_ids:
            selected_pid = st.selectbox("Select Your Paper ID", my_paper_ids)
            invited_eth = st.text_input("Invitee Ethereum Address")
            access_type = st.selectbox("Access Level", ["View", "Comment", "Edit"])
            access_map = {"View": 1, "Comment": 2, "Edit": 3}

            if invited_eth and st.button("Grant Access"):
                try:
                    tx = contract.functions.grantAccess(
                        int(selected_pid),
                        Web3.to_checksum_address(invited_eth),
                        access_map[access_type]
                    ).build_transaction({
                        'from': eth_address,
                        'nonce': w3.eth.get_transaction_count(eth_address),
                        'gas': 300000,
                        'gasPrice': w3.to_wei('50', 'gwei')
                    })
                    signed_tx = w3.eth.account.sign_transaction(tx, private_key)
                    tx_hash = w3.eth.send_raw_transaction(signed_tx.raw_transaction)
                    receipt = w3.eth.wait_for_transaction_receipt(tx_hash)
                    if receipt.status == 1:
                        st.success(f"✅ Access granted to {invited_eth} with {access_type} access.")
                    else:
                        st.error("❌ Failed to grant access.")
                except Exception as e:
                    st.error(f"Error: {str(e)}")
        else:
            st.info("You have not submitted any papers yet.")
    except Exception as e:
        st.warning(f"Grant access section failed: {str(e)}")
    # ==== Logs ====
    st.subheader("📜 Activity Logs")

    with st.expander("👥 User Registrations"):
        for e in get_events("UserRegistered"):
            block = w3.eth.get_block(e['blockNumber'])
            ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(block['timestamp']))
            name = e['args']['name']
            role = e['args']['role']
            role_str = {1: "Author", 2: "Faculty", 3: "Reviewer"}.get(role, "Unknown")
            st.info(f"[{ts}] Registered {name} as {role_str}")

    with st.expander("📄 Paper Submissions"):
        for e in get_events("PaperSubmitted"):
            block = w3.eth.get_block(e['blockNumber'])
            ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(block['timestamp']))
            author = e['args']['author']
            title = e['args']['title']
            st.info(f"[{ts}] '{title}' submitted by {author}")

    with st.expander("💬 Reviews"):
        for e in get_events("ReviewAdded"):
            block = w3.eth.get_block(e['blockNumber'])
            ts = time.strftime('%Y-%m-%d %H:%M:%S', time.localtime(block['timestamp']))
            reviewer = e['args']['reviewer']
            comment = e['args']['comment']
            st.info(f"[{ts}] Review by {reviewer}: {comment}")
else:
    st.info("👋 Welcome to ScholarChain! View or review papers as per your role.")
